"""
JeetMantra local-AI sidecar (FastAPI)
=====================================
A SMALL, OPTIONAL service the Node backend proxies to. It fills the three gaps
the codebase audit found, without touching the working Node/RAG spine:

  1. /extract  — pull text out of DOCX, images (OCR) and PDFs that rag.js skips
                 (rag.js:40 only handles PDF/TXT/MD/JSON/HTML). The returned text
                 feeds the EXISTING indexContent() chunker — nothing else changes.
  2. /embed    — local, free, multilingual embeddings (sentence-transformers)
                 so vectors don't depend on an OpenAI key. Write these into the
                 existing rag_chunks.embedding column.
  3. /tutor    — a local LLM answer via Ollama (offline-capable on a server/desktop).
                 Use as another provider behind aiProvider.js, or call directly.

This is a SERVER/desktop service. A genuinely "offline LLM per student in the
browser" is not realistic (a usable 7-8B model needs ~6-8 GB RAM/VRAM). Run this
where the hardware is; the browser stays thin and keeps its keyword offline
fallback (offlineTutorAnswer()).

Run:
    pip install -r requirements.txt
    uvicorn main:app --port 8000
Then set LOCAL_AI_URL=http://localhost:8000 in the Node backend (.env) and add
the proxy route shown in README.md.
"""
import io
import os
from typing import List, Optional

import httpx
from fastapi import FastAPI, UploadFile, File, HTTPException
from pydantic import BaseModel

app = FastAPI(title="JeetMantra Local AI", version="0.1.0")

# Lazy singletons — heavy imports/model loads happen on first use, not at boot,
# so the service starts instantly and only pays for what you call.
_embedder = None
OLLAMA_URL = os.environ.get("OLLAMA_URL", "http://localhost:11434")
OLLAMA_MODEL = os.environ.get("OLLAMA_MODEL", "llama3.1:8b")
EMBED_MODEL = os.environ.get("EMBED_MODEL", "BAAI/bge-m3")  # multilingual, 1024-dim


def _get_embedder():
    global _embedder
    if _embedder is None:
        from sentence_transformers import SentenceTransformer  # heavy import
        _embedder = SentenceTransformer(EMBED_MODEL)
    return _embedder


# ── text extraction ────────────────────────────────────────────────────────
def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((p.extract_text() or "") for p in reader.pages)


def _extract_docx(data: bytes) -> str:
    import docx  # python-docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_image(data: bytes) -> str:
    from PIL import Image
    import pytesseract
    return pytesseract.image_to_string(Image.open(io.BytesIO(data)))


@app.get("/health")
def health():
    return {"status": "ok", "embed_model": EMBED_MODEL, "ollama_model": OLLAMA_MODEL}


@app.post("/extract")
async def extract(file: UploadFile = File(...)):
    """DOCX / image-OCR / PDF / plain text → { text }. Feed into indexContent()."""
    data = await file.read()
    name = (file.filename or "").lower()
    ctype = (file.content_type or "").lower()
    try:
        if name.endswith(".pdf") or "pdf" in ctype:
            text = _extract_pdf(data)
        elif name.endswith(".docx") or "word" in ctype or "officedocument" in ctype:
            text = _extract_docx(data)
        elif any(name.endswith(e) for e in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff")) or ctype.startswith("image/"):
            text = _extract_image(data)
        else:
            text = data.decode("utf-8", errors="ignore")
    except Exception as e:  # never 500 the upload pipeline — return empty + reason
        return {"text": "", "chars": 0, "error": str(e)}
    text = (text or "").strip()
    return {"text": text, "chars": len(text)}


# ── local embeddings ───────────────────────────────────────────────────────
class EmbedIn(BaseModel):
    texts: List[str]


@app.post("/embed")
def embed(body: EmbedIn):
    """Multilingual local embeddings. Returns { vectors, dim }."""
    texts = [t for t in (body.texts or []) if isinstance(t, str) and t.strip()][:256]
    if not texts:
        return {"vectors": [], "dim": 0}
    vecs = _get_embedder().encode(texts, normalize_embeddings=True).tolist()
    return {"vectors": vecs, "dim": len(vecs[0]) if vecs else 0}


# ── local LLM tutor (Ollama) ───────────────────────────────────────────────
class TutorIn(BaseModel):
    prompt: str
    context: Optional[str] = None        # RAG chunks already retrieved by Node
    system: Optional[str] = None
    language: Optional[str] = "en"       # reply language (BCP-47 or short code)


@app.post("/tutor")
async def tutor(body: TutorIn):
    """Offline-capable answer via Ollama. Node sends already-retrieved context."""
    if not body.prompt:
        raise HTTPException(400, "prompt required")
    system = body.system or (
        "You are a patient tutor for school/college students. Answer using ONLY the "
        "provided context when it is relevant; if it does not contain the answer, say so. "
        f"Reply in language code '{body.language}'."
    )
    user = body.prompt
    if body.context:
        user = f"Context:\n{body.context}\n\nQuestion: {body.prompt}"
    payload = {
        "model": OLLAMA_MODEL,
        "system": system,
        "prompt": user,
        "stream": False,
    }
    try:
        async with httpx.AsyncClient(timeout=120) as client:
            r = await client.post(f"{OLLAMA_URL}/api/generate", json=payload)
            r.raise_for_status()
            data = r.json()
        return {"text": data.get("response", "").strip(), "model": OLLAMA_MODEL, "local": True}
    except Exception as e:
        # Let Node fall back to its cloud providers (aiProvider.js) on failure.
        raise HTTPException(502, f"local LLM unavailable: {e}")
